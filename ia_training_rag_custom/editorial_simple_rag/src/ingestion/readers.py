"""
Ingestion readers — one function per supported format.
Each reader returns: list of dicts with keys {text, source, metadata}.
"""

import json
import logging
import re
import warnings
from pathlib import Path

# Suppress spurious BeautifulSoup warning when a field value looks like a URL
# (e.g. a post's `link` field passed through _strip_html as a no-op).
try:
    from bs4 import MarkupResemblesLocatorWarning
    warnings.filterwarnings("ignore", category=MarkupResemblesLocatorWarning)
except ImportError:
    pass

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# JSON / CMS — helpers
# ---------------------------------------------------------------------------

def _strip_html(raw: str) -> str:
    """
    Strip HTML tags and decode entities.
    Uses BeautifulSoup when available; falls back to a regex stripper.
    """
    if not raw:
        return ""
    try:
        from bs4 import BeautifulSoup
        text = BeautifulSoup(raw, "html.parser").get_text(separator="\n", strip=True)
    except ImportError:
        # Crude fallback — replaces tags with spaces, collapses whitespace
        text = re.sub(r"<[^>]+>", " ", raw)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = text.strip()
    return text


def _rendered(value) -> str:
    """
    Extract text from a WordPress 'rendered' field.
    Accepts either a plain string (flat CMS format) or a dict
    with a 'rendered' key (WordPress REST API format).
    """
    if isinstance(value, dict):
        return value.get("rendered", "")
    return value or ""


# ---------------------------------------------------------------------------
# JSON / CMS
# ---------------------------------------------------------------------------

def read_json_cms(file_path: str) -> list[dict]:
    """
    Read a CMS JSON export (array of post objects).

    Supports two formats transparently:

    1. WordPress REST API format — title and content are objects:
       {
         "id": 13148,
         "date": "2026-02-06T11:34:55",
         "link": "https://example.com/?p=13148",
         "title":   {"rendered": "Post title here"},
         "content": {"rendered": "<p>HTML content…</p>", "protected": false}
       }

    2. Flat CMS format (legacy) — title and content are plain strings:
       {
         "id": "post-001",
         "title": "Post title here",
         "content": "Plain text content…",
         "url": "https://example.com/post-001",
         "published_at": "2026-03-15T10:30:00Z",
         "author": "…", "tags": […], "categories": […], "language": "en"
       }

    HTML tags in rendered fields are stripped via BeautifulSoup (or a regex
    fallback). The title is prepended to the content so retrieval can match
    on article headings.

    Returns list of dicts: {text, source, metadata}.
    """
    with open(file_path, encoding="utf-8") as f:
        posts = json.load(f)

    if not isinstance(posts, list):
        raise ValueError(f"{file_path}: expected a JSON array, got {type(posts)}")

    documents = []
    for post in posts:
        post_id = str(post.get("id", "?"))

        # --- title ---
        title_text = _strip_html(_rendered(post.get("title", "")))

        # --- content ---
        content_text = _strip_html(_rendered(post.get("content", "")))

        # Combine: title as first line so it is part of the retrieved chunk
        if title_text:
            text = title_text + "\n\n" + content_text
        else:
            text = content_text

        text = text.strip()
        if not text:
            logger.warning("Post %s has empty content — skipped", post_id)
            continue

        # --- source URL ---
        # WordPress: "link" field; flat format: "url" field; fallback: id
        source = post.get("link") or post.get("url") or post_id

        # --- date ---
        # WordPress: "date" field; flat format: "published_at"; fallback: ""
        published_at = post.get("date") or post.get("published_at", "")

        documents.append({
            "text": text,
            "source": source,
            "metadata": {
                "doc_type": "editorial_article",
                "post_id": post_id,
                "title": title_text,
                "author": post.get("author", ""),
                "published_at": published_at,
                "tags": post.get("tags", []),
                "categories": post.get("categories", []),
                "language": post.get("language", "en"),
                "source_file": str(file_path),
            },
        })

    logger.info("read_json_cms: %d documents from %s", len(documents), file_path)
    return documents


# ---------------------------------------------------------------------------
# CSV corpus (filename, text columns)
# ---------------------------------------------------------------------------

def read_csv_corpus(file_path: str) -> list[dict]:
    """
    Read a CSV corpus file with at minimum 'filename' and 'text' columns.
    Each row becomes one document. Used for journalist/archive corpora exported
    as flat CSV dumps (e.g. EPS_FILES_20K_NOV2025.txt).
    """
    import csv
    import sys
    csv.field_size_limit(sys.maxsize)
    documents = []
    with open(file_path, encoding="utf-8", newline="", errors="replace") as f:
        reader = csv.DictReader(f)
        for row in reader:
            filename = (row.get("filename") or "").strip()
            text = (row.get("text") or "").strip()
            if not text:
                continue
            documents.append({
                "text": text,
                "source": filename or str(file_path),
                "metadata": {
                    "doc_type": "journalist_article",
                    "source_file": str(file_path),
                    "filename": filename,
                    "language": "",
                },
            })
    logger.info("read_csv_corpus: %d documents from %s", len(documents), file_path)
    return documents


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------

def read_txt(file_path: str) -> list[dict]:
    with open(file_path, encoding="utf-8", errors="replace") as f:
        first_line = f.readline().strip()
        f.seek(0)
        # Detect CSV corpus format: header line is exactly "filename,text"
        if first_line == "filename,text":
            return read_csv_corpus(file_path)
        text = f.read().strip()
    if not text:
        return []
    return [{
        "text": text,
        "source": str(file_path),
        "metadata": {
            "doc_type": "txt",
            "source_file": str(file_path),
            "language": "",
        },
    }]


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def read_pdf(file_path: str) -> list[dict]:
    try:
        import pypdfium2 as pdfium
    except ImportError:
        raise ImportError("pypdfium2 is required for PDF ingestion: pip install pypdfium2")

    pdf = pdfium.PdfDocument(file_path)
    documents = []
    for i, page in enumerate(pdf):
        text_page = page.get_textpage()
        text = (text_page.get_text_range() or "").strip()
        if not text:
            continue
        documents.append({
            "text": text,
            "source": f"{file_path}#page={i + 1}",
            "metadata": {
                "doc_type": "pdf",
                "page": i + 1,
                "source_file": str(file_path),
                "language": "",
            },
        })
    return documents


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def read_docx(file_path: str) -> list[dict]:
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx is required for DOCX ingestion: pip install python-docx")

    doc = docx.Document(file_path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if not text:
        return []
    return [{
        "text": text,
        "source": str(file_path),
        "metadata": {
            "doc_type": "docx",
            "source_file": str(file_path),
            "language": "",
        },
    }]


# ---------------------------------------------------------------------------
# TSV / XLSX
# ---------------------------------------------------------------------------

def read_tsv(file_path: str) -> list[dict]:
    try:
        import csv
    except ImportError:
        pass  # csv is stdlib

    rows = []
    with open(file_path, encoding="utf-8", newline="") as f:
        reader = csv.DictReader(f, delimiter="\t")
        for row in reader:
            rows.append(row)

    text = "\n".join(str(r) for r in rows)
    if not text:
        return []
    return [{
        "text": text,
        "source": str(file_path),
        "metadata": {
            "doc_type": "tsv",
            "source_file": str(file_path),
            "language": "",
        },
    }]


def read_xlsx(file_path: str) -> list[dict]:
    try:
        import openpyxl
    except ImportError:
        raise ImportError("openpyxl is required for XLSX ingestion: pip install openpyxl")

    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    rows = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            row_text = "\t".join(str(c) if c is not None else "" for c in row)
            if row_text.strip():
                rows.append(row_text)

    text = "\n".join(rows)
    if not text:
        return []
    return [{
        "text": text,
        "source": str(file_path),
        "metadata": {
            "doc_type": "xlsx",
            "source_file": str(file_path),
            "language": "",
        },
    }]


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

READERS = {
    ".json": read_json_cms,
    ".txt":  read_txt,
    ".pdf":  read_pdf,
    ".docx": read_docx,
    ".tsv":  read_tsv,
    ".xlsx": read_xlsx,
}


def read_file(file_path: str) -> list[dict]:
    """Dispatch to the right reader based on file extension."""
    ext = Path(file_path).suffix.lower()
    reader = READERS.get(ext)
    if reader is None:
        logger.warning("No reader for extension %s — skipping %s", ext, file_path)
        return []
    return reader(file_path)
